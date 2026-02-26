import sys
from pathlib import Path
sys.path.append(str(Path(__file__).resolve().parent.parent))

import os
from pypdf import PdfReader
from sentence_transformers import SentenceTransformer
from qdrant_client import QdrantClient
from qdrant_client.models import VectorParams, Distance, PointStruct
from docx import Document
import openpyxl
import time

from config import EMBED_MODEL_PATH, DOCS_DIR, QDRANT_PATH, COLLECTION_NAME


# =====================================
# TEXT CHUNKING
# =====================================

def chunk_text(text, max_length=800, overlap=150):
    chunks = []
    start = 0

    while start < len(text):
        end = start + max_length
        chunk = text[start:end]
        chunks.append(chunk.strip())
        start += max_length - overlap

    return chunks


# =====================================
# LOADERS
# =====================================

def load_txt(path):
    return path.read_text(encoding="utf-8")


def load_pdf(path):
    reader = PdfReader(str(path))
    text = ""
    for page in reader.pages:
        extracted = page.extract_text()
        if extracted:
            text += extracted + "\n"
    return text


def load_docx(path):
    doc = Document(str(path))
    full_text = []

    # параграфы
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            full_text.append(paragraph.text)

    # таблицы
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells)
            if row_text.strip():
                full_text.append(row_text)

    return "\n".join(full_text)


def load_xlsx(path):
    wb = openpyxl.load_workbook(path, data_only=True)
    documents = []

    for sheet in wb.worksheets:
        headers = None
        row_index = 0

        for row in sheet.iter_rows(values_only=True):
            row_index += 1

            if not headers:
                headers = row
                continue

            row_dict = {
                str(headers[i]): str(row[i])
                for i in range(len(headers))
                if headers[i] and row[i] is not None
            }

            if row_dict:
                row_text = " | ".join(
                    f"{k}: {v}" for k, v in row_dict.items()
                )

                documents.append({
                    "text": f"Лист: {sheet.title} | {row_text}",
                    "sheet": sheet.title,
                    "row": row_index,
                })

    return documents

def load_document(path):
    filename = path.name.lower()

    if filename.endswith(".txt"):
        return load_txt(path)

    elif filename.endswith(".pdf"):
        return load_pdf(path)

    elif filename.endswith(".docx"):
        return load_docx(path)

    elif filename.endswith(".xlsx"):
        return load_xlsx(path)

    else:
        return None


# =====================================
# MAIN INDEXING
# =====================================

def main():
    model = SentenceTransformer(str(EMBED_MODEL_PATH))
    client = None

    try:
        client = QdrantClient(path=str(QDRANT_PATH))
    except RuntimeError as e:
        message = str(e)
        if "already accessed by another instance" in message:
            print(
                "\n[INDEX ERROR] Папка qdrant_storage занята другим процессом.\n"
                "Остановите приложение/процесс, использующий Qdrant, и повторите команду:\n"
                "python -m scripts.index_to_qdrant\n"
            )
            return
        raise

    points = []
    point_id = 0
    BATCH_SIZE = 256

    # Пересоздаём коллекцию
    client.recreate_collection(
        collection_name=COLLECTION_NAME,
        vectors_config=VectorParams(
            size=1024,
            distance=Distance.COSINE
        ),
    )

    for filename in os.listdir(DOCS_DIR):

        path = DOCS_DIR / filename
        if not path.is_file():
            continue

        content = load_document(path)
        if not content:
            continue

        print(f"Обработка файла: {filename}")

        # ==================================================
        # EXCEL → ROW-BASED INDEXING
        # ==================================================
        if filename.lower().endswith(".xlsx"):

            # content теперь список строк-документов
            for row_doc in content:

                embedding = model.encode(
                    row_doc["text"],
                    normalize_embeddings=True
                )

                points.append(
                    PointStruct(
                        id=point_id,
                        vector=embedding.tolist(),
                        payload={
                            "text": row_doc["text"],
                            "source": filename,
                            "sheet": row_doc["sheet"],
                            "row": row_doc["row"],
                            "doc_type": "excel_row"
                        }
                    )
                )

                if len(points) >= BATCH_SIZE:
                    client.upsert(
                        collection_name=COLLECTION_NAME,
                        points=points
                    )
                    points = []

                point_id += 1

        # ==================================================
        # TXT / PDF / DOCX → SEMANTIC CHUNKING
        # ==================================================
        else:

            chunks = chunk_text(content)

            if not chunks:
                continue

            embeddings = model.encode(
                chunks,
                normalize_embeddings=True
            )

            for chunk, embedding in zip(chunks, embeddings):

                points.append(
                    PointStruct(
                        id=point_id,
                        vector=embedding.tolist(),
                        payload={
                            "text": chunk,
                            "source": filename,
                            "doc_type": "text_chunk"
                        }
                    )
                )
                if len(points) >= BATCH_SIZE:
                    client.upsert(
                        collection_name=COLLECTION_NAME,
                        points=points
                    )
                    points = []
                point_id += 1

    print(f"\nВсего проиндексировано объектов: {point_id}")

    if points:
        client.upsert(
            collection_name=COLLECTION_NAME,
            points=points
        )
    if point_id > 0:
        print("Индексация в Qdrant завершена.")
    else:
        print("Нет данных для индексации.")

    if client is not None:
        # Небольшая пауза перед закрытием, чтобы избежать гонок на Windows lock
        time.sleep(0.2)
        client.close()


if __name__ == "__main__":
    main()
