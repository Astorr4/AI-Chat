import uuid
import asyncio
import time
import logging
import json
from collections import defaultdict, deque
import getpass
from io import BytesIO

from fastapi import FastAPI, Request, UploadFile, File
from fastapi.staticfiles import StaticFiles
from fastapi.responses import FileResponse, StreamingResponse, JSONResponse
from pydantic import BaseModel
from pypdf import PdfReader
from docx import Document
import openpyxl
from core.embeddings import EmbeddingModel
from core.vector_store import VectorStore
from core.llm import LLM
from core.rag import RAG
from app.services.chat_service import ChatService
from app.services.stats_service import build_rag_stats
from core.database import (
    init_db,
    save_session,
    save_message,
    increment_question_stat,
    cleanup_old_data,
    save_rag_metrics,
)
from config import (
    RATE_LIMIT,
    RATE_WINDOW,
    LLM_CONCURRENCY_LIMIT,
)

# =====================================
# CONFIG
# =====================================

llm_semaphore = asyncio.Semaphore(LLM_CONCURRENCY_LIMIT)


# =====================================
# LOGGING
# =====================================

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s | %(levelname)s | %(message)s",
)

logger = logging.getLogger("AI-Assistant")

# =====================================
# APP INIT
# =====================================

app = FastAPI(title="Corporate AI Assistant")
app.mount("/static", StaticFiles(directory="app/static"), name="static")
init_db()

# =====================================
# RAG INIT
# =====================================

embed = EmbeddingModel()
store = VectorStore()
llm = LLM()
rag = RAG(embed, store, llm)
chat_service = ChatService(rag=rag, llm_semaphore=llm_semaphore, logger=logger)

# =====================================
# RATE LIMIT
# =====================================

rate_limit_store = defaultdict(deque)
file_analysis_contexts = {}
MAX_FILE_CONTEXT_CHARS = 50000

def check_rate_limit(session_id):
    now = time.time()
    window_start = now - RATE_WINDOW

    timestamps = rate_limit_store[session_id]

    while timestamps and timestamps[0] < window_start:
        timestamps.popleft()

    if len(timestamps) >= RATE_LIMIT:
        return False

    timestamps.append(now)
    return True

# =====================================
# CLEANUP TASK
# =====================================

async def cleanup_task():
    while True:
        cleanup_old_data(
            days_messages=7,
            days_metrics=30,
            days_question_stats=180
        )
        await asyncio.sleep(600)

@app.on_event("startup")
async def startup_event():
    asyncio.create_task(cleanup_task())


# =====================================
# MODELS
# =====================================

class Query(BaseModel):
    question: str


def _extract_text_from_pdf(raw_bytes: bytes):
    reader = PdfReader(BytesIO(raw_bytes))
    parts = []
    for page in reader.pages:
        extracted = page.extract_text()
        if extracted:
            parts.append(extracted)
    return "\n".join(parts)


def _extract_text_from_docx(raw_bytes: bytes):
    document = Document(BytesIO(raw_bytes))
    parts = []

    for paragraph in document.paragraphs:
        text = (paragraph.text or "").strip()
        if text:
            parts.append(text)

    for table in document.tables:
        for row in table.rows:
            row_text = " | ".join((cell.text or "").strip() for cell in row.cells)
            if row_text.strip():
                parts.append(row_text)

    return "\n".join(parts)


def _extract_text_from_xlsx(raw_bytes: bytes):
    workbook = openpyxl.load_workbook(filename=BytesIO(raw_bytes), data_only=True)
    rows = []

    for sheet in workbook.worksheets:
        headers = None
        for row in sheet.iter_rows(values_only=True):
            if headers is None:
                headers = list(row)
                continue

            row_map = {}
            for idx, value in enumerate(row):
                if value is None or idx >= len(headers):
                    continue
                header = headers[idx]
                if header is None:
                    continue
                row_map[str(header)] = str(value)

            if row_map:
                payload = " | ".join(f"{k}: {v}" for k, v in row_map.items())
                rows.append(f"Лист: {sheet.title} | {payload}")

    return "\n".join(rows)


def _extract_text_from_uploaded_file(filename: str, raw_bytes: bytes):
    lowered = (filename or "").lower()

    if lowered.endswith(".pdf"):
        return _extract_text_from_pdf(raw_bytes), "pdf"
    if lowered.endswith(".docx"):
        return _extract_text_from_docx(raw_bytes), "docx"
    if lowered.endswith(".xlsx"):
        return _extract_text_from_xlsx(raw_bytes), "xlsx"

    return None, None


# =====================================
# ROUTES
# =====================================

@app.get("/")
def serve_ui():
    return FileResponse("app/static/index.html")


@app.post("/chat")
async def chat(query: Query, request: Request):
    user_login = getpass.getuser()
    session_id = request.headers.get("X-Session-Id")
    if not session_id:
        session_id = str(uuid.uuid4())

    if not check_rate_limit(session_id):
        return JSONResponse(
            status_code=429,
            content={"error": "Превышен лимит запросов."},
        )

    save_session(session_id)
    save_message(session_id, "user", query.question)
    increment_question_stat(query.question)
    request_id, stream = chat_service.stream_chat(
        question=query.question,
        session_id=session_id,
        user_login=user_login,
    )

    if request_id is None:
        return JSONResponse(
            status_code=503,
            content={"error": "Сервис перегружен, попробуйте чуть позже."},
        )

    response = StreamingResponse(
        stream,
        media_type="text/plain",
    )

    response.headers["X-Session-Id"] = session_id
    response.headers["X-Request-Id"] = request_id
    return response


@app.post("/file-analysis/upload")
async def upload_file_for_analysis(request: Request, file: UploadFile = File(...)):
    session_id = request.headers.get("X-Session-Id")
    if not session_id:
        session_id = str(uuid.uuid4())

    if not file or not file.filename:
        return JSONResponse(status_code=400, content={"error": "Файл не передан."})

    raw_bytes = await file.read()
    if not raw_bytes:
        return JSONResponse(status_code=400, content={"error": "Файл пустой."})

    text, file_type = _extract_text_from_uploaded_file(file.filename, raw_bytes)
    if not text or not text.strip() or not file_type:
        return JSONResponse(
            status_code=400,
            content={"error": "Поддерживаются только файлы .pdf, .docx, .xlsx"},
        )

    normalized_text = text.strip()
    is_truncated = False
    if len(normalized_text) > MAX_FILE_CONTEXT_CHARS:
        normalized_text = normalized_text[:MAX_FILE_CONTEXT_CHARS]
        is_truncated = True

    file_analysis_contexts[session_id] = {
        "file_name": file.filename,
        "file_type": file_type,
        "context": normalized_text,
        "uploaded_at": time.time(),
    }

    save_session(session_id)

    response = JSONResponse(
        {
            "session_id": session_id,
            "file_name": file.filename,
            "file_type": file_type,
            "context_chars": len(normalized_text),
            "truncated": is_truncated,
        }
    )
    response.headers["X-Session-Id"] = session_id
    return response


@app.post("/file-analysis/chat")
async def file_analysis_chat(query: Query, request: Request):
    session_id = request.headers.get("X-Session-Id")
    if not session_id:
        return JSONResponse(status_code=400, content={"error": "Отсутствует X-Session-Id."})

    file_ctx = file_analysis_contexts.get(session_id)
    if not file_ctx:
        return JSONResponse(
            status_code=400,
            content={"error": "Сначала загрузите файл через кнопку «Анализ файла»."},
        )

    save_session(session_id)
    save_message(session_id, "user", query.question)
    increment_question_stat(query.question)

    messages = [
        {
            "role": "system",
            "content": (
                "Ты анализируешь загруженный пользователем документ. "
                "Отвечай строго по содержимому документа. "
                "Разрешены: пересказ, ответы на конкретные вопросы, своды по таблицам, "
                "определение лучших/худших показателей. "
                "Если данных нет в документе, прямо сообщи об этом."
            ),
        },
        {
            "role": "user",
            "content": (
                f"Документ: {file_ctx['file_name']}\n"
                f"Содержимое:\n{file_ctx['context']}\n\n"
                f"Вопрос пользователя:\n{query.question}"
            ),
        },
    ]

    # Generate request ID to support streaming recovery
    request_id, stream = chat_service.stream_chat(
        question=query.question,
        session_id=session_id,
        user_login="file_analysis",
        custom_messages=messages,
        is_file_analysis=True
    )

    if request_id is None:
        return JSONResponse(
            status_code=503,
            content={"error": "Сервис перегружен, попробуйте чуть позже."},
        )

    response = StreamingResponse(
        stream,
        media_type="text/plain",
    )

    response.headers["X-Session-Id"] = session_id
    response.headers["X-Request-Id"] = request_id
    return response


@app.get("/file-analysis-stream/{request_id}")
async def file_analysis_stream(request_id: str, request: Request):
    session_id = request.headers.get("X-Session-Id")
    if not session_id:
        return JSONResponse(
            status_code=400,
            content={"error": "Отсутствует X-Session-Id."},
        )

    allowed, reason = chat_service.can_access_generation(request_id, session_id)
    if not allowed:
        if reason == "not_found":
            return JSONResponse(status_code=404, content={"error": "Поток не найден."})
        return JSONResponse(status_code=403, content={"error": "Доступ запрещён."})

    return StreamingResponse(
        chat_service.stream_generation(request_id),
        media_type="text/plain",
    )


@app.get("/chat-stream/{request_id}")
async def chat_stream(request_id: str, request: Request):
    session_id = request.headers.get("X-Session-Id")
    if not session_id:
        return JSONResponse(
            status_code=400,
            content={"error": "Отсутствует X-Session-Id."},
        )

    allowed, reason = chat_service.can_access_generation(request_id, session_id)
    if not allowed:
        if reason == "not_found":
            return JSONResponse(status_code=404, content={"error": "Поток не найден."})
        return JSONResponse(status_code=403, content={"error": "Доступ запрещён."})

    return StreamingResponse(
        chat_service.stream_generation(request_id),
        media_type="text/plain",
    )
@app.post("/chat-sync")
async def chat_sync(query: Query, request: Request):

    session_id = request.headers.get("X-Session-Id")
    if not session_id:
        session_id = str(uuid.uuid4())

    save_session(session_id)
    save_message(session_id, "user", query.question)

    retrieval_start = time.time()

    cleaned_question = rag.normalize_query(query.question)

    result = rag.search(cleaned_question, top_k=3, session_id=session_id)

    documents = result["documents"]
    similarity = result.get("similarity")
    topic_mode = result.get("topic_mode")

    if not documents:
        return {
            "answer": "В документации информация не найдена.",
            "confidence": 0,
            "rejected": True
        }

    scores = [
        d.get("hybrid_score", d.get("score", 0))
        for d in documents
    ]

    max_score = max(scores)
    avg_score = sum(scores) / len(scores)

    context = "\n".join(d["text"] for d in documents)

    messages = [
        {
            "role": "system",
            "content": "Отвечай строго по контексту."
        },
        {
            "role": "user",
            "content": f"Контекст:\n{context}\n\nВопрос:\n{query.question}"
        }
    ]

    answer = rag.llm.generate(messages)

    final_confidence = rag.recalibrate_confidence(
        max_score=max_score,
        avg_score=avg_score,
        coverage=1.0,
        similarity=similarity,
        topic_mode=topic_mode
    )
    save_rag_metrics(
        session_id=session_id,
        user_login="load_test",
        question=query.question,
        found_docs=len(documents),
        filtered_docs=len(documents),
        confidence=final_confidence,
        avg_score=avg_score,
        min_score=min(scores),
        max_score=max_score,
        coverage=1.0,
        threshold=0.0,
        sources_count=len(documents),
        context_chars=len(context),
        retrieval_time=round(time.time() - retrieval_start, 3),
        llm_time=0,
        total_time=round(time.time() - retrieval_start, 3),
        answer_length=len(answer),
        is_followup=0,
        memory_size=0,
        rejected_reason=None,
        similarity=similarity,
        topic_mode=topic_mode
    )
    return {
        "answer": answer,
        "confidence": final_confidence,
        "rejected": False
    }
@app.get("/rag-stats")
def rag_stats(from_date: str = None, to_date: str = None):
    data = build_rag_stats(from_date=from_date, to_date=to_date)
    return JSONResponse(data)


@app.get("/stats")
def serve_stats():
    return FileResponse("app/static/stats.html")
